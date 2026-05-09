#!/usr/bin/env python3
"""Generate Word report for Garbage Detection using Surveillance Camera project."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent / "Garbage_Detection_Surveillance_Report.docx"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None, header_bg="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "D6E4F0" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def heading(doc, text, level=1, color="1F4E79"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )
    return p


def body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if p.runs:
        p.runs[0].font.size = Pt(10.5)
    return p


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Garbage Detection Using Surveillance Cameras", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.size = Pt(18)

sub = doc.add_paragraph("Capstone Project Report — Model Comparison & Results")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
sub.runs[0].font.size = Pt(12)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OBJECTIVE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Objective")
body(doc, (
    "The goal of this project is to build an automated garbage detection system "
    "for surveillance cameras (CCTV). The system must detect the presence and location "
    "of garbage in video frames using object detection models, enabling smart city "
    "infrastructure to identify littering without manual monitoring. "
    "The system targets real-time or near-real-time inference with high recall — "
    "missing a garbage instance is more costly than a false alarm in a surveillance context."
))


# ══════════════════════════════════════════════════════════════════════════════
# 2. DATASETS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Datasets")

heading(doc, "2.1  Dataset 1 — Garbage Detection using CCTV (COCO format)", level=2, color="2E75B6")
body(doc, (
    "Source: Roboflow — COCO annotation format. All images were captured from fixed-angle "
    "surveillance cameras and contain a single object class: garbage."
))

add_table(doc,
    ["Property", "Value"],
    [
        ["Total Images",       "3,861"],
        ["Total Annotations",  "10,948"],
        ["Object Class",       "garbage (single class)"],
        ["Annotation Format",  "COCO JSON"],
        ["Image Resolution",   "640 × 640 px"],
        ["Train Split",        "3,051 images / 8,638 annotations (80%)"],
        ["Val Split",          "381 images / 1,158 annotations (10%)"],
        ["Test Split",         "382 images / 1,152 annotations (10%)"],
    ],
    col_widths=[2.0, 4.0]
)
doc.add_paragraph()

heading(doc, "2.2  Dataset 2 — Garbage Detection using CCTV v2 (YOLO format)", level=2, color="2E75B6")
body(doc, (
    "A second, larger dataset sourced from Roboflow in YOLO format. "
    "Same single class (garbage), same surveillance camera domain. "
    "This dataset was later merged with Dataset 1 to increase training data."
))

add_table(doc,
    ["Property", "Value"],
    [
        ["Train Images",  "9,152"],
        ["Val Images",    "779"],
        ["Object Class",  "garbage (single class)"],
        ["Format",        "YOLO (images + labels)"],
    ],
    col_widths=[2.0, 4.0]
)
doc.add_paragraph()

heading(doc, "2.3  Merged Dataset", level=2, color="2E75B6")
body(doc, (
    "Dataset 1 (COCO) was converted to YOLO format and merged with Dataset 2. "
    "Filename collisions were avoided by prefixing Dataset 2 images with 'ds2_'. "
    "The test set was kept as Dataset 1 only to ensure evaluation integrity — "
    "the model is tested on unseen data from the original distribution."
))

add_table(doc,
    ["Split", "Dataset 1", "Dataset 2", "Merged Total"],
    [
        ["Train", "3,051", "9,152", "12,203"],
        ["Val",   "381",   "779",   "1,160"],
        ["Test",  "382",   "—",     "382 (unchanged)"],
    ],
    col_widths=[1.2, 1.5, 1.5, 1.8]
)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 3. MODELS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Models Used")

heading(doc, "3.1  Florence-2 (Microsoft)", level=2, color="2E75B6")
body(doc, (
    "Florence-2-large is a 780M parameter vision-language foundation model that performs "
    "object detection by generating text sequences containing location tokens "
    "(e.g., garbage<loc_234><loc_112><loc_456><loc_389>). "
    "It was evaluated in two modes:"
))
bullet(doc, "Zero-Shot: No training. Florence-2 is prompted with <OD> and predicted boxes are filtered for garbage-related labels.")
bullet(doc, "Fine-Tuned: LoRA adapter (rank=16) trained on Dataset 1 for 10 epochs using the existing stage_2 fine-tuning pipeline. Only 0.45% of parameters (3.5M / 780M) were trainable.")

heading(doc, "3.2  YOLOv8m (Ultralytics)", level=2, color="2E75B6")
body(doc, (
    "YOLOv8 medium (25.8M parameters) is a purpose-built real-time object detector. "
    "It regresses bounding boxes directly using a decoupled detection head, making it "
    "significantly faster and more precise than generalist vision-language models on "
    "narrow single-class detection tasks. "
    "Multiple training configurations were evaluated to arrive at the best setup."
))


# ══════════════════════════════════════════════════════════════════════════════
# 4. EXPERIMENT 1 — DATASET 1 ONLY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Experiment 1 — Dataset 1 Only (3,814 images)")

body(doc, (
    "The first set of experiments used Dataset 1 alone. All three models were trained "
    "and evaluated on the same 80/10/10 split. Results are reported on the test set (382 images, 1,152 annotations)."
))

heading(doc, "4.1  Results — Dataset 1 Only", level=2, color="2E75B6")
add_table(doc,
    ["Model", "Precision", "Recall", "F1", "mAP@0.5", "mAP@0.5:0.95"],
    [
        ["Florence-2 Zero-Shot",  "0.216", "0.049", "0.081", "—",     "—"],
        ["Florence-2 Fine-Tuned", "0.109", "0.075", "0.089", "—",     "—"],
        ["YOLOv8m (v2, AdamW)",   "0.481", "0.366", "0.415", "0.424", "0.210"],
        ["YOLOv8m (v3, SGD)",     "0.486", "0.414", "0.447", "0.440", "0.210"],
    ],
    col_widths=[2.2, 1.0, 1.0, 1.0, 1.0, 1.3]
)
doc.add_paragraph()

heading(doc, "4.2  Why Florence-2 Was Dropped", level=2, color="2E75B6")
body(doc, "Florence-2 performed poorly on this task for three fundamental reasons:")
bullet(doc, "Architectural mismatch: Florence-2 generates bounding boxes as text tokens, which is inherently less precise than YOLO's direct coordinate regression.")
bullet(doc, "Generalist vs. specialist: Florence-2 is optimised for open-vocabulary, multi-task scenarios. A single-class CCTV task plays to YOLO's strengths.")
bullet(doc, "LoRA bottleneck: Only 0.45% of parameters were trained. The frozen backbone limits how much the model can adapt to the CCTV domain.")
bullet(doc, "Time cost vs. gain: Training Florence-2 for 30+ epochs would take 3+ hours and realistically only push F1 from 0.089 to ~0.15–0.20 — still far behind YOLO.")
body(doc, (
    "Conclusion: All further experiments focused exclusively on YOLOv8m, "
    "which is the architecturally correct tool for this narrow detection task."
))


# ══════════════════════════════════════════════════════════════════════════════
# 5. YOLO CONFIG EVOLUTION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. YOLOv8m Configuration Evolution")
body(doc, (
    "Three YOLO configurations were evaluated to understand the impact of "
    "hyperparameter choices on detection performance."
))

add_table(doc,
    ["Setting", "v2 (Initial)", "v3 (Fixed)", "v4 (Final)"],
    [
        ["Optimizer",       "AdamW",  "SGD",    "SGD"],
        ["Learning Rate",   "0.001",  "0.001",  "0.0005"],
        ["Image Size",      "1024",   "1024",   "1024"],
        ["copy_paste aug.", "0.3",    "0.1",    "0.1"],
        ["multi_scale",     "True",   "False",  "False"],
        ["flipud",          "0.0",    "0.0",    "0.0"],
        ["dropout",         "0.1",    "0.0",    "0.0"],
        ["Dataset",         "DS1",    "DS1",    "DS1 + DS2"],
        ["Best Epoch",      "94/134", "23/63",  "150/150"],
    ],
    col_widths=[1.8, 1.3, 1.3, 1.3]
)
doc.add_paragraph()

body(doc, "Key lessons from configuration tuning:", bold=True)
bullet(doc, "AdamW → SGD: SGD with momentum 0.937 proved more stable for this dataset, improving recall by +4.8 points.")
bullet(doc, "flipud=0.0: CCTV cameras have fixed orientation — flipping upside-down created unrealistic training samples.")
bullet(doc, "copy_paste reduced to 0.1: At 0.3, this augmentation introduced too much noise given the small batch size of 4.")
bullet(doc, "lr0=0.0005: The lower learning rate prevented early convergence (v3 peaked at epoch 23; v4 trained the full 150 epochs).")


# ══════════════════════════════════════════════════════════════════════════════
# 6. EXPERIMENT 2 — MERGED DATASET
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Experiment 2 — Merged Dataset (12,203 training images)")

body(doc, (
    "The primary bottleneck identified after Experiment 1 was dataset size. "
    "3,051 training images is insufficient to learn the full diversity of garbage "
    "appearances across different lighting conditions, distances, and occlusion levels. "
    "Dataset 2 (9,152 images, same class, same domain) was merged into the training and "
    "validation splits, quadrupling the available training data."
))

heading(doc, "6.1  Merge Process", level=2, color="2E75B6")
bullet(doc, "Dataset 1 was converted from COCO to YOLO format (split_cctv_coco.py → coco_to_yolo.py).")
bullet(doc, "Dataset 2 was already in YOLO format — images and labels were directly compatible (nc=1, class=garbage).")
bullet(doc, "Dataset 2 files were prefixed with 'ds2_' to prevent filename collisions.")
bullet(doc, "All 12,203 image-label pairs were verified: zero mismatches, all coordinates normalised in [0, 1].")
bullet(doc, "Test set was kept as Dataset 1 only (382 images) to ensure a clean, uncontaminated evaluation.")

heading(doc, "6.2  Results — Merged Dataset (YOLOv8m v4)", level=2, color="2E75B6")
add_table(doc,
    ["Metric", "Value"],
    [
        ["mAP@0.5",          "0.910"],
        ["mAP@0.5:0.95",     "0.727"],
        ["Precision",        "0.925"],
        ["Recall",           "0.850"],
        ["F1 Score",         "0.886"],
        ["Training Epochs",  "150 / 150 (no early stopping)"],
        ["Training Time",    "~13.3 hours on NVIDIA A100 40GB"],
    ],
    col_widths=[2.5, 3.5]
)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 7. FULL COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Complete Results Comparison")
body(doc, "All models evaluated on the same test set: 382 images, 1,152 garbage annotations.")

add_table(doc,
    ["Model", "Dataset", "Precision", "Recall", "F1", "mAP@0.5"],
    [
        ["Florence-2 Zero-Shot",      "DS1",        "0.216", "0.049", "0.081", "—"],
        ["Florence-2 Fine-Tuned",     "DS1",        "0.109", "0.075", "0.089", "—"],
        ["YOLOv8m v2 (AdamW)",        "DS1",        "0.481", "0.366", "0.415", "0.424"],
        ["YOLOv8m v3 (SGD)",          "DS1",        "0.486", "0.414", "0.447", "0.440"],
        ["YOLOv8m v4 (SGD+Merged)",   "DS1 + DS2",  "0.925", "0.850", "0.886", "0.910"],
    ],
    col_widths=[2.2, 1.2, 1.0, 1.0, 1.0, 1.0]
)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 8. ANALYSIS & CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Analysis & Conclusion")

body(doc, "Key Findings:", bold=True)
bullet(doc, "Data volume was the single biggest factor. Merging datasets pushed F1 from 0.447 → 0.886 — a 2x improvement that no hyperparameter change achieved.")
bullet(doc, "YOLOv8m is the right architecture for this task. Florence-2, despite being a state-of-the-art foundation model, achieved only F1=0.089 — 10x worse than YOLOv8m on this narrow single-class detection task.")
bullet(doc, "Configuration tuning had moderate impact (~7% F1 gain). The most impactful individual change was switching from AdamW to SGD (+recall) and lowering lr0 to 0.0005 (prevented early stopping).")
bullet(doc, "The final model (F1=0.886, Recall=0.850, mAP@0.5=0.910) meets deployment thresholds for a real surveillance system. It correctly detects 85% of garbage instances with 92.5% precision.")

body(doc, "Deployment Recommendation:", bold=True)
body(doc, (
    "The YOLOv8m model trained on the merged dataset is recommended for deployment. "
    "At a confidence threshold of 0.15, the model achieves the best balance of precision "
    "and recall for a surveillance context where missed detections are more costly than "
    "false alarms. Inference speed is 3.8ms per image on an A100 GPU, supporting real-time "
    "processing of CCTV streams at 25+ FPS."
))

body(doc, "Limitations & Future Work:", bold=True)
bullet(doc, "Night-time and low-light scenarios: HSV augmentation was applied but dedicated night-vision datasets would further improve robustness.")
bullet(doc, "Small distant objects: SAHI (Slicing Aided Hyper Inference) can boost recall for tiny garbage objects by slicing images into overlapping tiles at inference time.")
bullet(doc, "Florence-2 potential: With 30+ epochs and full fine-tuning (not LoRA), Florence-2 could potentially be competitive — but is not practical given compute constraints.")


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT_PATH))
print(f"Report saved: {OUT_PATH}")
